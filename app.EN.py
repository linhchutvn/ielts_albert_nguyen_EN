import streamlit as st
import google.generativeai as genai
import json
import re
import time
from PIL import Image
import random
import textwrap
import html
import os
import requests
from io import BytesIO

# Library for Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Library for PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# ==========================================
# 1. API & PROMPT CONFIGURATION
# ==========================================

ALL_KEYS = st.secrets["GEMINI_API_KEYS"]

def generate_content_with_failover(prompt, image=None):
    """Smart function to automatically detect the best available Model with quota."""
    keys_to_try = list(ALL_KEYS)
    random.shuffle(keys_to_try) 
    
    # PRIORITY LIST
    model_priority = [
        #"gemini-2.0-flash-thinking-preview-01-21",
        #"gemini-3-pro-preview", 
        #"gemini-2.5-pro",
        "gemini-3-flash-preview",        
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-2.0-flash",
        "gemini-1.5-pro", 
        "gemini-1.5-flash"
    ]
    
    last_error = ""
    for index, current_key in enumerate(keys_to_try):
        try:
            genai.configure(api_key=current_key)
            
            # Get list of models actually available for this key
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            
            # Find best model
            sel_model = None
            for target in model_priority:
                if any(target in m_name for m_name in available_models):
                    sel_model = target
                    break
            
            if not sel_model:
                sel_model = "gemini-1.5-flash" 

            # --- DISPLAY MODEL INFO ---
            masked_key = f"****{current_key[-4:]}"
            
            st.toast(f"⚡ Connected: {sel_model}", icon="🤖")
            
            with st.expander("🔌 Technical Connection Details (Debug)", expanded=False):
                st.write(f"**Active Model:** `{sel_model}`")
                st.write(f"**Active API Key:** `{masked_key}` (Key #{index + 1})")
                if "thinking" in sel_model.lower():
                    st.caption("🧠 Thinking Mode: ON")
            # ------------------------------------------------
            
            temp_model = genai.GenerativeModel(
                model_name=sel_model, 
            )
            
            content_parts = [prompt]
            if image:
                content_parts.append(image)
                
            # Generation Config
            gen_config = {
                "temperature": 0.3,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 32000,
            }

            if "thinking" in sel_model.lower():
                gen_config["thinking_config"] = {
                    "include_thoughts": True,
                    "thinking_budget": 32000
                }

            response = temp_model.generate_content(
                content_parts,
                generation_config=gen_config
            )
            
            return response, sel_model 
            
        except Exception as e:
            last_error = str(e)
            if "429" in last_error or "quota" in last_error.lower() or "limit" in last_error.lower():
                continue 
            else:
                break
                
    st.error(f"❌ All {len(keys_to_try)} Keys have exceeded their quota. Last error: {last_error}")
    return None, None 

# --- ENGLISH PROMPT TEMPLATE ---
GRADING_PROMPT_TEMPLATE = """

Please assume the role of an **IELTS Examiner with 30 years of experience at the British Council**, specializing in exam design and grading for IELTS Writing Task 1. Your mission is to grade the following submission based on the official Band 9.0 criteria with absolute strictness and precision.

**Exam Classification (Context Awareness):** You must correctly identify whether the submission is **IELTS Academic** (Charts, Graphs, Processes, or Maps) and apply the corresponding set of Band Descriptors accordingly.

**SPECIAL REQUIREMENT (DEEP SCRUTINY MODE):**
Do not provide a quick response. Take your time to "think" and perform a step-by-step, highly detailed analysis.

### 1. CORE WORKING PROTOCOL

*   **>> SLOW REASONING PROTOCOL:**
    *   You are strictly forbidden from summarizing your feedback. For each criterion, you must write at least 200–300 words.
    *   Perform a **"Socratic Analysis"**: Question every sentence written by the candidate, identify every imperfection, and explain exhaustively why it fails to reach Band 7.0 or Band 9.0 based on the specific data in the text.
    *   Generic phrases such as "Good grammar" or "Appropriate vocabulary" are prohibited. You must cite **at least 3–5 specific examples** from the essay for every criterion to substantiate your judgment.

*   **Persona:** You are a veteran examiner—stern and demanding, yet fair. Your tone must be direct and clinical. Do not use hollow flattery. If the essay is poor, state so clearly.

*   **>> HOLISTIC SCORING PRINCIPLE:** You must strictly distinguish between a **Systematic Error** and a **Slip**.
    *   *Slip:* A minor, random error (e.g., a single missing letter or one-off comparative error). If the essay otherwise demonstrates superior linguistic control, these slips **MUST NOT** be used as a justification to downgrade a score from 8 to 7 or from 9 to 8.

*   **"Deep Scan" Mode:** Do not rush. Spend time analyzing every sentence and every word through a "Step-by-Step Analysis" workflow.

*   **Exhaustive Listing Rule:**
    *   Absolutely **DO NOT** group errors. If the candidate makes 10 article errors, you must list all 10 items individually.
    *   The error list in the JSON output serves as "legal evidence." Every minor error (commas, capitalization, articles) must be recorded. An empty or sparse JSON list combined with a low GRA score is considered a major logical contradiction.
    *   **>> TAXONOMY RULE:** When categorizing errors in JSON, use only standard linguistic and examiner terminology (e.g., *Subject-Verb Agreement, Collocation, Article, Comma Splice*). Do not invent non-standard terms like "Bad word" or "Wrong grammar."

*   **>> TWO-PASS SCANNING:**
    *   *Pass 1:* Identify macro errors (Sentence structure, misused academic register, data logic, and task fulfillment).
    *   *Pass 2:* Re-scan the entire text for micro errors (Articles, singular/plural agreement, punctuation, and capitalization). 
    *   The final error list must only be compiled after completing both passes.

*   **>> APPROXIMATION TOLERANCE:**
    *   For very small figures (e.g., < 2-3%), accept strong approximation language such as *"virtually no"*, *"almost zero"*, or *"negligible"*. Do not mark these as data inaccuracies (Logic Errors) unless the actual figure exceeds 5%.

### 2. DETAILED GRADING CRITERIA (4 CRITERIA)

#### A. Task Achievement (TA)
*   **Data Reasoning & Information Grouping:**
    *   **Band 8.0+:** Candidates MUST demonstrate skillful selection and logical grouping of similar data points within paragraphs. Mechanical listing will be capped at **Band 6.0-7.0**.
    *   **>> ADDED COMPARISON RULE:** If the report only provides a linear description of data without establishing correlations or comparisons between objects -> **MAX BAND 6.0** (even if data is 100% accurate).
    *   **>> ADDED "TOTAL/OTHER" SAFETY NET:** Categories such as 'Total', 'Miscellaneous', or 'Other' are NOT mandatory key features. No points shall be deducted if the candidate chooses to omit them.
*   **Word Count & Conciseness:**
    *   **No Unfair Penalty:** Reports > 200 words with high-value information and 100% accuracy shall NOT have TA scores lowered.
    *   **Penalty conditions:** Only deduct marks if the writing is wordy due to repetition or irrelevance. For high-quality reports > 220 words, provide a "Tip" regarding conciseness rather than a score deduction.
    *   **Penalties:** < 150 words (strict TA evaluation); < 20 words (Band 1).
*   **"Fatal" Negative Features (TA):**
    *   **Object vs. Figure:** Harshly penalize subject-object confusion (e.g., "The figure of apples rose" is INCORRECT; "The consumption of apples rose" is CORRECT).
    *   **Wrong Units:** Confusing percentages (%) with whole numbers caps TA at **Band 5.0**.
    *   **No Data/Support:** Academic reports describing trends without supporting figures = **Band 5.0**.
    *   **Band 5 (Critical):** If trends are described without supporting data, the score MUST be lowered to **Band 5.0** per the bolded descriptor: *"There may be no data to support the description."*
    *   **Overview Requirements:** Processes must cover Start-Middle-End; Maps must show the overall transformation. Missing/Incorrect Overview = **Max Band 5.0-6.0**. 
    *   **Band 7:** Must identify clear main trends or differences (Clear overview).
    *   **Band 6:** Some effort to provide an overview, but information may be poorly selected or unclear.
    *   **Band 5:** No overview or the overview is completely inaccurate.
    *   **Personal Opinion:** Strictly prohibited. Inclusion of personal views results in a heavy penalty.
*   **>> ADDED FORMAT & TONE RULES:**
    *   **Format Error:** Using bullet points or numbered lists instead of paragraphs = **MAX BAND 5.0 TA**.
    *   **Tone Error (GT):** Using informal language (slang, contractions like "gonna") in a "Formal letter" = Penalty down to **Band 5.0-6.0**.
*   **Math Logic Check:** Scrutinize adverbs of degree (e.g., *slight* vs. *significant*). Example: An increase from 10% to 15% is a 50% relative increase; therefore, using "slight" is logically incorrect.
*   **Endpoint Trap:** Strictly forbid the use of the word "peak" for the final data point on a graph (as the future trend is unknown). Suggest: "reaching a high of."
*   **>> OVERVIEW STRATEGY (BAND 8.0-9.0):**
    1.  **"No Data" Principle:** High-band overviews MUST NOT contain specific figures.
    2.  **Double Content Structure:** Must cover both (1) Main Trends AND (2) Major Comparisons/High-lows.
    3.  **Synthesis Technique:** Evaluate whether the candidate synthesizes similar objects or simply lists them.
    4.  **Placement:** Encourage placement immediately after the Introduction for optimal logical flow.

#### B. Coherence & Cohesion (CC)
*   **Invisible Cohesion (Band 9):** Prioritize structures like "respectively", "in that order", and reduced relative clauses.
*   **Mechanical Linkers:** Over-reliance on "Firstly, Secondly, In addition, Furthermore" at the start of every sentence = **Max Band 6.0**.
*   **Paragraphing:** Must be logical. A single-block essay = **Max Band 5.0 CC**.
*   **>> ADDED "AMBIGUOUS REFERENCING" (The 'It' Trap):** 
    *   Strictly check pronouns (It, This, That, These, Those). If the antecedent is unclear, causing reader confusion = **MAX BAND 6.0 CC**.
*   **>> ADDED "INVISIBLE GLUE" RULE:**
    *   Scrutinize signposting words. Starting paragraphs with "Regarding..." or "As for..." more than twice is marked as **Mechanical (Band 6.0/7.0)**.
    *   Encourage transitions via sentence subjects or referencing (e.g., instead of "Regarding A, it increased...", use "A, conversely, witnessed a rise...").
*   **>> CC FLEXIBILITY PRINCIPLE:** If logic and clarity are high, slightly mechanical linkers should not automatically drop the score to 7.0. Aim for **Band 8.0** if the flow is smooth. Only drop to 7.0 if linkers are disruptive.
*   **>> OUTPUT REQUIREMENTS:** 
    *   **Evidence-based:** Must quote specific sentences from the candidate's work for analysis.
    *   **Adaptive Suggestions:** 
        *   Below Band 7: Suggest fixes for ACCURACY.
        *   Band 7+: Suggest upgrades for NATURALNESS (Band 9 style).

#### C. Lexical Resource (LR)
*   **Naturalness over Academic:** Prefer natural vocabulary (use, help, start) over pretentious or misused academic jargon (utilise, facilitate, commence).
*   **Blacklist:** Flag clichéd/memorized formulaic language.
*   **Precision:** Evaluate collocations (e.g., "increased significantly" is better than "increased strongly").
*   **>> ADDED REPETITION RULE:** 
    *   Repeating key vocabulary (e.g., "increase", "fluctuate") > 3 times without attempting to paraphrase = **MAX BAND 5.0 LR** (Limited flexibility).
*   **>> SPELLING THRESHOLD:**
    *   1-2 minor slips = Potential Band 8.0.
    *   A few errors (readable) = Band 7.0.
    *   Noticeable errors = Band 6.0.
    *   Meaning impeded = Band 5.0.
*   **>> NO DOUBLE PENALIZATION PRINCIPLE:** 
    *   Spelling and Redundancy errors should be penalized under LR, not GRA, provided the sentence structure remains intact. A candidate can still achieve **9.0 GRA** with minor lexical slips.
*   **Word Choice:** Prefer "Proportion" or "Share" for workforce/population data; "Percentage" is a raw figure. "Chosen one" is marked as informal/inappropriate for economic contexts.

#### D. Grammatical Range & Accuracy (GRA)
*   **Absolute Accuracy:** Scrutinize articles, prepositions, and singular/plural agreement.
*   **Error-free Sentence Ratio:**
    *   Band 6: Errors present but meaning is clear.
    *   Band 7: Error-free sentences are frequent.
    *   Band 8+: The majority of sentences are completely error-free.
*   **Technical Errors:**
    *   **Comma Splice:** Joining independent clauses with only a comma = Drops score to **Band 5.0-6.0**.
    *   **The Mad Max:** Overuse or omission of the definite article "the".
    *   **Past Perfect Trigger:** "By + [past time]" requires the Past Perfect tense. Failure to use it indicates poor range.
*   **>> ADDED PUNCTUATION CONTROL:** Beyond Comma Splices, frequent lack of commas in subordinate clauses or arbitrary capitalization = **Capped below Band 8.0 GRA**.
*   **>> PARAPHRASING STRATEGY (Intro):** 
    *   Identify the opening sentence. Converting a Noun Phrase (the number of...) into a Noun Clause (how many...) is a hallmark of **Band 8.0+ GRA**.
*   **Band 9 Threshold:** If the writing uses natural, complex structures, allow 1-2 minor slips. Do not cap at 8.0 for a single article error.
*   **>> "SLIPS" PRINCIPLE:** Band 9.0 GRA allows for "rare minor errors." If the candidate uses a wide range of complex structures naturally, do not hesitate to award a 9.0 despite 1 or 2 slips. Avoid mechanical capping at 8.0.

### 3. SCORING PROCESS & SELF-CORRECTION PROTOCOL (STRICT 1:1 SYNC)

**CORE MANDATE:** Every single word or punctuation mark enclosed within `<del>...</del>` tags in the revised essay **MUST** have a corresponding, individual entry in the `errors` list. Summarizing or merging multiple errors into a single entry is strictly prohibited.

**Step 1: Deep Scan & Error Documentation (JSON Errors Array)**
*   Perform a 3-pass scan of the essay and list **ALL** identified issues in the `errors` array.
*   **>> MANDATORY EVIDENCE RULE:** 
    *   If you assign a **Coherence & Cohesion (CC) score below 9.0**, you are **REQUIRED** to create at least **2-3 specific error entries** in the `errors` array under the `Coherence & Cohesion` category to justify the penalty. 
    *   *Example:* If CC is 6.0, you must explicitly document issues such as: "Paragraph 2 lacks a clear topic sentence," "The linker 'Moreover' is used incorrectly," or "Logical flow is disrupted."
    *   **PROHIBITED:** You must never leave the CC error list empty if the CC score is lower than 9.0.
*   **Two-Pass Execution Detail:**
    *   *Pass 1 (Grammar/Vocab):* Meticulously inspect every article, comma, and singular/plural usage.
    *   *Pass 2 (Data Logic):* Verify "Object vs. Figure" logic (e.g., identifying if the candidate mistakenly used "industry" as the subject instead of "industrial emissions").
*   **Full Enumeration:** Populate the `errors` array first. If there are 14 incorrect instances in the text, there must be exactly 14 error objects in the JSON. 
    *   *Example:* If the article "the" is missing in 3 different locations, you must create 3 separate error entries.
*   **>> DOUBLE-TAGGING RULE (NEW):**
    *   If you encounter a severe grammatical error that also disrupts the logical flow (e.g., `Sentence Fragment`, `Run-on Sentence`, `Comma Splice`), you must create **TWO** separate error entries:
        1.  A `Grammar` entry (to correct the syntax).
        2.  A `Coherence & Cohesion` entry with the error type `Fragmented Flow` (to penalize the lack of coherence).
    *   This ensures the CC section is populated and prevents the system from displaying inaccurate "Excellent" feedback when structural issues exist.
*   Calculate the Band Scores for the original essay (Markdown).
*   **Rounding Rule:** Remainder of .125 -> round down to .0; .25 -> round UP to .5; .375 -> round UP to .5; .625 -> round down to .5; .75 -> round UP to the next whole number.

**Step 2: Annotated Essay Generation**
*   **Mirroring Principle:** You are only permitted to correct errors that were explicitly listed in the JSON `errors` array in Step 1.
*   **No Hidden Edits:** Strictly forbid "silent fixes" (such as fixing capitalization or adding a missing "the") within the annotated essay if those errors were not officially declared in the `errors` list.
*   The total count of `<del>` tags **MUST** exactly equal the number of entries in the `errors` list. Any discrepancy will be treated as a serious protocol violation.

**Step 3: Internal Re-grading (JSON revised_score)**
*   Assume the role of an independent second examiner to grade the `annotated_essay` as if it were a fresh submission (with micro-errors fixed).
*   **Content Rule:** Since this revision primarily fixes GRA/LR while maintaining the original structure, the Task Achievement (TA) and Coherence & Cohesion (CC) scores **SHOULD GENERALLY REMAIN THE SAME** as the original. If the original essay lacks an Overview or contains data inaccuracies, the revised score must reflect these persistent flaws.
*   **Revised Score Constraints:**
    *   **Word Count Check:** If the revision exceeds 200 words, TA is capped at **8.0** (penalty for lack of conciseness/economy).
    *   **Naturalness Check:** If pretentious or overly academic vocabulary is used inappropriately, LR is capped at **8.0**.
*   **Consistency & Parity Check:** 
    *   Count the `<del>` tags in the revision. If they do not match the number of entries in the `errors` array (e.g., 14 edits but only 7 declared errors), you have failed the protocol. You must re-generate the JSON `errors` array to achieve a **1:1 ratio**.
*   **>> THE 9.0 BARRIER:**
    *   **Coherence & Cohesion (CC):** Strictly **DO NOT** award a 9.0 if the structure still relies on mechanical linkers at the start of sentences (e.g., "Regarding...", "In addition...", "Overall..."). Band 9 CC requires "invisible cohesion." If the original structure is at a Band 7-8 level, the revised CC score **MUST** stay at 7-8.
    *   **Task Achievement & Lexical (TA/LR):** Re-verify "Object vs. Figure" logic. If the candidate wrote "Industry was the most polluted" instead of "Industrial emissions were the highest," this is a fundamental data logic error. Even if grammar is corrected, TA and LR must be capped at **7.0 - 8.0**.
    *   **Unit Accuracy:** Scrutinize units (tonnes, %, numbers). If the original confused units, the revised TA score cannot increase by more than 1.0 band.
*   **>> FINAL RE-SCAN PROTOCOL:** Before finalizing the `revised_score`, ask yourself: *"Am I being too generous? Does this revision still possess the 'skeleton' of a Band 7 essay?"* If so, lower the score immediately to ensure examiner stringency.

### INFORMATION:
a/ Task Prompt: {{TOPIC}}
b/ Visual Data Note: {{IMAGE_NOTE}}
c/ Candidate's Report: {{ESSAY}}

---
### DETAILED EVALUATION CONTENT:

**CRITICAL PEDAGOGY RULE:**
When providing correction examples or rewrites, you must align them with the **current Band score** of the submission:
*   **If the score is < 6.0:** Provide a rewrite at **Band 7.0** level (Focus on Accuracy, Clarity, and Simplicity). Avoid overly complex jargon.
*   **If the score is >= 6.5:** Provide a rewrite at **Band 9.0** level (Focus on Sophistication, Academic Register, and Complex Syntactical Structures).

**ANTI-BREVITY RULES:**
1.  **Strict Prohibition of Generic Comments:** Do not write vague feedback such as "Improve your grammar." You must specify the exact area (e.g., tenses, articles, or sentence structure).
2.  **Mandatory Citation of Evidence:** Every observation must be supported by quoting specific sentences or phrases directly from the candidate's text.
3.  **Mandatory Modeling:** Regardless of whether the essay is Band 1 or Band 9, you **MUST** provide rewrite examples at the end of every criterion section. This is non-negotiable.

---

### **1. Task Achievement (TA):**

*   **Overview Assessment:** [Analyze the overview: Is there one? Is it placed optimally? Does it capture the main trends and major comparisons? *Note: Band 9 requires a sophisticated overview that generalizes rather than just listing data.*]
*   **Warning for Band 5-6:** [If the overview contains detailed figures/data, explain why this traps the candidate at Band 5. Instruct them on how to remove data to reach Band 7.]
*   **Accuracy and Data Selection:** [Verify data accuracy. Is there "Data Saturation" (listing too many trivial figures)? **Reminder: Ignore 'Total'/'Other' categories when assessing completeness if they are not significant.**]
*   **Response Strategy:** [Evaluate the information grouping. Is the candidate describing data linearly (Band 5 style) or using logical synthesis to compare and contrast (Band 7+ style)?]

*   **⚠️ Critical Errors & In-depth Analysis:** 
    *   [For every error found, you **MUST** explain it using the following 3 steps:
        1. **Quote the error:** (e.g., "the figure of pizza ate")
        2. **Linguistic reason:** (e.g., "Selectional Restriction Violation" or "Object vs. Figure logic error").
        3. **Impact:** (e.g., "Confuses the reader regarding the subject, diminishing the academic tone").]

*   **💡 BAND UPGRADE STRATEGY (STEP-BY-STEP):**
    *   **Step 1 (Filter):** Strictly remove data from the overview. Focus on the "meaning" of the numbers.
    *   **Step 2 (Synthesize):** Group objects with similar trends to ensure conciseness (Economy).
    *   **Step 3 (Contrast):** Always highlight the highest/lowest points or significant rank changes.
    *   **Step 4 (Link):** Use "Invisible Cohesion" (While/Whereas/V-ing) instead of mechanical linkers.

*   **✍️ MODEL COMPARISON (CHOOSE THE APPROPRIATE LEVEL):**
    *   **Realistic Model (Target Band 7.0):** 
        *   *"This is a clear, accurate version that you can achieve immediately by refining your current logic:"*
        *   **[AI: PROVIDE A BAND 7.0 OVERVIEW & BODY SAMPLE BASED ON CANDIDATE'S IDEAS]**
    *   **Advanced Model (Reference Band 9.0):** 
        *   *"This is a native-level version for your reference, demonstrating sophisticated vocabulary and data synthesis:"*
        *   **[AI: PROVIDE A BAND 9.0 OVERVIEW & BODY SAMPLE HERE]**

> **📍 Task Achievement Score:** [Score/9.0]

---

#### **2. Coherence and Cohesion (CC):**

*   **Paragraphing Logic:** [Analyze the grouping: Is it based on Time, Object, or Trend? Does this help the reader compare data easily? Does each paragraph have a clear focal point?]
*   **Linking Devices:** [Evaluate naturalness:
    *   **Warning:** Is there an over-reliance on sentence-initial "Mechanical Linking" (e.g., *Regarding, Turning to, Firstly*)?
    *   **Encouragement:** Is "Invisible Cohesion" used (e.g., mid-sentence adverbs like *meanwhile, however* or relative clauses)?]
*   **Referencing:** [Check referencing techniques: Are *it, this, that, the former, the latter, respectively* used correctly to avoid repetition?]
*   **⚠️ Specific Weaknesses:** [Identify:
    1.  **Fragmented Flow:** Isolated sentences.
    2.  **Ambiguous Referencing:** Unclear antecedents for pronouns.
    3.  **Repetitive Sentence Openers:** Starting every sentence with "The figure...".
    4.  **Sentence Fragments:** Missing main verbs.]
*   **💡 Correction & Upgrade:**
    *   *Candidate’s Original (Issue):* "[Quote exact phrase]"
    *   *Proposed Rewrite (Natural Flow):* "[If Band < 7: Fix for ACCURACY. If Band 7+: Rewrite using advanced cohesive structures for Band 9]."
    *   *Explanation:* "[Why is the new version more professional?]"
*   **Mandatory Depth Requirement:** For every error, follow the 3-step explanation: 
    1. Quote error.
    2. Descriptor-based reason.
    3. Impact on communication.

> **📍 Coherence & Cohesion Score:** [Score/9.0]

---

#### **3. Lexical Resource (LR):**

*   **Range & Flexibility Assessment:** [Is the vocabulary basic, adequate, or sophisticated? Is there "Repetition" of keywords (e.g., increase, decrease, figure)?]
*   **Precision & Style:** [Are collocations natural? Is there "Word-for-word translation" from the mother tongue? Is the register too informal (e.g., "get up" instead of "increase")?]
*   **⚠️ Core Lexical Weaknesses:** [Do not just list spelling errors. Identify **systemic habits**, e.g., "You frequently misuse economic terminology" or "You use pretentious language inappropriately."]
*   **💡 Vocabulary Upgrade:**
    *   *Repetitive word used:* "[e.g., 'increase']"
    *   *Suggested replacements:* 
        *   *[For Band < 7]:* Basic but accurate (rise, growth, climb).
        *   *[For Band 7+]:* Academic/Sophisticated (escalate, upsurge, register a growth).
*   **Mandatory Depth Requirement:** For every error, follow the 3-step explanation: 
    1. Quote error. 
    2. Descriptor-based reason. 
    3. Impact on communication.

> **📍 Lexical Resource Score:** [Score/9.0]

---

#### **4. Grammatical Range and Accuracy (GRA):**

*   **Range Analysis:** [Does the writing rely on simple/compound sentences? Are there Band 8+ structures like *Passive Voice, Reduced Relative Clauses, or Nominalization*?]
*   **Accuracy Check:** [Estimate the **Error-free sentence ratio**: Below 50% (Band 5), 50-70% (Band 6-7), or > 80% (Band 8+)? Distinguish between **Systematic Errors** and **Slips**. *Note: If a single minor slip is the only error, maintain Band 8.5-9.0.*]
*   **⚠️ Systematic Errors to Fix:** [Identify the candidate's biggest grammatical gap (e.g., articles, tenses, or complex clause coordination).]
*   **💡 Sentence Transformation Challenge:**
    *   *Original Sentence:* "[Quote a simple or erroneous sentence]"
    *   *Upgraded Version:* 
        *   *[For Band < 7]:* Combine into a clear complex sentence using *because, although, or which*.
        *   *[For Band 7+]:* Transform using advanced grammar (Inversion, Participle Phrases, or Nominalization).
*   **Mandatory Depth Requirement:** For every error, follow the 3-step explanation: 
    1. Quote error. 
    2. Descriptor-based reason. 
    3. Impact on communication.

> **📍 Grammatical Range & Accuracy Score:** [Score/9.0]

---

### **OVERALL BAND SCORE:** [Apply the .25/.75 rounding rule]

---

### **EXAMINER'S STRATEGIC TIPS:**
1.  **Strategic Advice:** Provide tips based on actual patterns observed in the essay.
2.  **Economy:** How to prune redundant words (especially if the essay is > 200 words).
3.  **Introduction Power:** Demonstrate how to convert a Noun Phrase into a Noun Clause in the introduction to boost GRA.
4.  **Grouping:** How to group data more intelligently (e.g., Highs vs. Lows).
5.  **Overview Mastery:** Final specific advice on crafting a high-band overview.

#### **5. ANALYSIS DATA (JSON):**

Must extract data into a single **JSON Object**.

**ALLOWED ERROR TYPES (TAXONOMY):**

**A. [COHERENCE & COHESION] - Macro Errors:**
`Illogical Grouping`, `Missing Overview`, `Fragmented Flow`, `Lack of Progression`, `Incoherent Paragraphing`, `Mechanical Linking`, `Overuse of Connectors`, `Ambiguous Referencing`, `Repetitive Structure`, `Data Inaccuracy`.

**B. [GRAMMAR] - Micro Errors:**
`Comma Splice`, `Run-on Sentence`, `Sentence Fragment`, `Faulty Parallelism`, `Misplaced Modifier`, `Word Order`, `Subject-Verb Agreement`, `Tense Inconsistency`, `Passive Voice Error`, `Relative Clause Error`, `Article Error`, `Preposition Error`, `Singular/Plural`, `Countable/Uncountable`, `Punctuation`.

**C. [VOCABULARY] - Lexical Errors:**
`Imprecise Word Choice`, `Incompatible Collocation`, `Word Form Error`, `Selectional Restriction Violation`, `Informal Register`, `Pretentious Language`, `Redundancy`, `Forced Paraphrasing`.

**INTERNAL RE-GRADING OF REVISED ESSAY (MOST CRITICAL STEP):**
   - Forget that you just corrected this essay. Assume the role of a second, independent Examiner grading the newly generated 'annotated_essay'.
   - **Content Rule:** The revised version only corrects Grammar/Vocabulary; it CANNOT fix original errors related to missing data or a lack of comparison. If the original TA was 6.0, the revised TA remains at 6.0 (or at most 7.0 if clarity is significantly improved).
   - **Conclusion:** The 'revised_score' MUST be the actual score of the revised essay; it MUST NOT default to 9.0.

JSON Structure:
```json
{
  "original_score": {
      "task_achievement": "TA score of the original essay (User's essay)",
      "cohesion_coherence": "CC score of the original essay",
      "lexical_resource": "LR score of the original essay",
      "grammatical_range": "GRA score of the original essay",
      "overall": "Overall score of the original essay (Average)"
  },
  "errors": [
    {
      "category": "Grammar" or "Vocabulary",
      "type": "Error Type",
      "impact_level": "High" | "Medium" | "Low",
      "explanation": "Brief explanation of the error.",
      "original": "the incorrect text snippet",
      "correction": "the correct text snippet (IN ALL CAPS)"
    }
  ],
  "annotated_essay": "The revised version of the essay (maintaining original paragraph structure). Wrap incorrect words in <del>...</del> tags and corrected words in <ins class='grammar'>...</ins> or <ins class='vocab'>...</ins> tags. The corrected content must be IN ALL CAPS.",
   "revised_score": {
      "word_count_check": "MANDATORY: STATE THE WORD COUNT OF THE REVISED ESSAY (e.g., '235 words - Too long')",
      "logic_re_evaluation": "Explain any score deductions (e.g., 'Despite being grammatically flawless, the essay is 235 words long, violating the principle of conciseness, thus TA is capped at 8.0').",
      "task_achievement": "The actual TA score (penalize heavily for wordiness)",
      "cohesion_coherence": "CC score",
      "lexical_resource": "LR score",
      "grammatical_range": "GRA score",
      "overall": "Average score (rounded according to IELTS rules)"
  }
}
```
"""

# ==========================================
# 2. UI CONFIGURATION
# ==========================================
st.set_page_config(page_title="IELTS Examiner Pro", page_icon="🛡️", layout="wide")

# CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@300;400;700&display=swap');
    
    /* Font size for Report */
    .report-content {
        font-size: 19px !important;
        line-height: 1.7 !important;
        color: #1F2937;
    }
    .report-content ul, .report-content ol {
        margin-bottom: 15px;
    }
    .report-content li {
        margin-bottom: 8px;
    }
    .report-content strong {
        color: #0F172A;
        font-weight: 700;
    }

    /* Global Fonts */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    /* Header Style */
    h1 {
        font-family: 'Merriweather', serif !important;
        color: #0F172A !important;
        font-weight: 700 !important;
    }
    .pro-badge {
        color: #D40E14; 
        font-weight: bold;
    }
    .verified-badge {
        background-color: #F1F5F9;
        border: 1px solid #E2E8F0;
        padding: 4px 12px;
        border-radius: 99px;
        font-size: 14px;
        font-weight: bold;
        color: #475569;
        display: inline-flex;
        align-items: center;
        margin-left: 10px;
    }
    
    /* Error Cards */
    .error-card {
        background-color: white;
        border: 1px solid #E5E7EB;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        transition: all 0.2s;
    }
    .error-card:hover {
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        border-color: #D1D5DB;
    }
    .error-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 12px;
        border-bottom: 1px solid #F3F4F6;
        padding-bottom: 8px;
    }
    .error-badge-grammar {
        background-color: #DCFCE7;
        border: 1px solid #22C55E;
        color: #022C22;
        padding: 2px 8px;
        border-radius: 6px;
        font-size: 14px;
        font-weight: 800;
        text-transform: uppercase;
    }
    .error-badge-vocab {
        background-color: #FEF9C3;
        border: 1px solid #FCD34D;
        color: #713F12;
        padding: 2px 8px;
        border-radius: 6px;
        font-size: 14px;
        font-weight: 800;
        text-transform: uppercase;
    }
    .impact-high { background-color: #FEE2E2; color: #991B1B; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #FECACA;}
    .impact-medium { background-color: #FFEDD5; color: #9A3412; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #FED7AA;}
    .impact-low { background-color: #DBEAFE; color: #1E40AF; padding: 2px 8px; border-radius: 99px; font-size: 14px; font-weight: bold; border: 1px solid #BFDBFE;}
    
    .correction-box {
        background-color: #F9FAFB;
        padding: 12px;
        border-radius: 8px;
        margin-bottom: 12px;
        font-size: 16px;
        border: 1px solid #F3F4F6;
    }
    
    /* Annotated Essay Style */
    .annotated-text {
        font-family: 'Merriweather', serif;
        line-height: 1.8;
        color: #374151;
        background-color: white;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
    }
    del {
        color: #9CA3AF;
        text-decoration: line-through;
        margin-right: 4px;
        text-decoration-thickness: 2px;
    }
    ins.grammar {
        background-color: #4ADE80;
        color: #022C22;
        text-decoration: none;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 700;
        border: 1px solid #22C55E;
    }
    ins.vocab {
        background-color: #FDE047;
        color: #000;
        text-decoration: none;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 700;
        border: 1px solid #FCD34D;
    }
    
    /* Button Style */
    div.stButton > button {
        background-color: #D40E14;
        color: white;
        font-weight: bold;
        border: none;
        padding: 10px 24px;
        border-radius: 8px;
        transition: all 0.3s;
    }
    div.stButton > button:hover {
        background-color: #B91C1C;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. AI CONNECTION & DATA PROCESSING
# ==========================================

def clean_json(text):
    """Extract JSON from AI response"""
    match = re.search(r"```json\s*([\s\S]*?)\s*```", text)
    if match:
        content = match.group(1)
        content = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', content)
        return content.strip()
    return None

def calculate_overall(scores):
    """Calculate IELTS Overall Score"""
    try:
        valid_scores = []
        for s in scores:
            try:
                valid_scores.append(float(s))
            except:
                continue
                
        if not valid_scores or len(valid_scores) < 4: return '-'
        
        avg = sum(valid_scores) / len(valid_scores)
        decimal = avg - int(avg)
        
        # IELTS Rounding Rules
        if decimal < 0.25: final = int(avg)
        elif decimal < 0.75: final = int(avg) + 0.5
        else: final = int(avg) + 1.0
        
        return str(final)
    except:
        return '-'

def process_response(full_text):
    """Process AI Response: Separate Markdown and JSON"""
    json_str = clean_json(full_text)
    markdown_part = full_text
    
    data = {
        "errors": [], 
        "annotatedEssay": None, 
        "revisedScore": None, 
        "originalScore": {
            "task_achievement": "-",
            "cohesion_coherence": "-",
            "lexical_resource": "-",
            "grammatical_range": "-",
            "overall": "-"
        }
    }
    
    # A. Parse JSON
    if json_str:
        markdown_part = full_text.split("```json")[0].strip()
        try:
            parsed = json.loads(json_str)
            data["errors"] = parsed.get("errors", [])
            data["annotatedEssay"] = parsed.get("annotated_essay")
            data["revisedScore"] = parsed.get("revised_score")
        except json.JSONDecodeError:
            pass

    # B. Extract Scores via Regex (Updated for English)
    patterns = {
        "task_achievement": r"Task\s+Achievement\s*Score.*?(\d+\.?\d*)",
        "cohesion_coherence": r"Coherence\s*&\s*Cohesion\s*Score.*?(\d+\.?\d*)",
        "lexical_resource": r"Lexical\s+Resource\s*Score.*?(\d+\.?\d*)",
        "grammatical_range": r"Grammatical\s+Range.*?Score.*?(\d+\.?\d*)",
    }
    
    found_scores = []
    
    for key, regex in patterns.items():
        match = re.search(regex, markdown_part, re.IGNORECASE | re.DOTALL)
        if match:
            score = match.group(1)
            data["originalScore"][key] = score
            found_scores.append(score)
        else:
            try:
                if json_str:
                    parsed = json.loads(json_str)
                    val = parsed.get("original_score", {}).get(key, "-")
                    data["originalScore"][key] = str(val)
                    if str(val) != "-": found_scores.append(val)
            except:
                pass

    if found_scores:
        data["originalScore"]["overall"] = calculate_overall(found_scores)

    return markdown_part, data

# --- FILE EXPORT FUNCTIONS ---

def register_fonts():
    """Download and register Roboto font"""
    font_reg = "Roboto-Regular.ttf"
    font_bold = "Roboto-Bold.ttf"
    
    urls = {
        font_reg: "https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Regular.ttf",
        font_bold: "https://github.com/googlefonts/roboto/raw/main/src/hinted/Roboto-Bold.ttf"
    }
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }

    try:
        for filename, url in urls.items():
            if os.path.exists(filename) and os.path.getsize(filename) < 1000:
                os.remove(filename) 
                
            if not os.path.exists(filename):
                response = requests.get(url, headers=headers, timeout=20)
                if response.status_code == 200:
                    with open(filename, "wb") as f:
                        f.write(response.content)
                else:
                    print(f"❌ Failed to download {filename}. Status: {response.status_code}")

        if os.path.exists(font_reg) and os.path.exists(font_bold):
            pdfmetrics.registerFont(TTFont('Roboto', font_reg))
            pdfmetrics.registerFont(TTFont('Roboto-Bold', font_bold))
            addMapping('Roboto', 0, 0, 'Roboto') 
            addMapping('Roboto', 1, 0, 'Roboto-Bold')
            return True
        else:
            return False
            
    except Exception as e:
        print(f"❌ Font Error: {e}")
        return False

# WORD EXPORT
def create_docx(data, topic, original_essay, analysis_text):
    doc = Document()
    
    heading = doc.add_heading('IELTS WRITING TASK 1 - ASSESSMENT REPORT', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {time.strftime('%d/%m/%Y')}")
    
    # 1. BAND SCORE
    doc.add_heading('1. BAND SCORE', level=1)
    scores = data.get("originalScore")
    
    if scores and isinstance(scores, dict) and scores.get('overall', '-') != '-':
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        
        headers = ['Task Achievement', 'Coherence', 'Lexical Resource', 'Grammar', 'OVERALL']
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        
        vals = [
            str(scores.get('task_achievement', '-')),
            str(scores.get('cohesion_coherence', '-')),
            str(scores.get('lexical_resource', '-')),
            str(scores.get('grammatical_range', '-')),
            str(scores.get('overall', '-'))
        ]
        for i, v in enumerate(vals):
            table.cell(1, i).text = v
    else:
        doc.add_paragraph("Score details could not be extracted automatically.")

    # 2. ANALYSIS
    doc.add_heading('2. EXAMINER\'S DETAILED ANALYSIS', level=1)
    if analysis_text:
        clean_analysis = analysis_text.replace('**', '').replace('### ', '').replace('#### ', '')
        doc.add_paragraph(clean_analysis)

    # 3. ERRORS
    doc.add_heading('3. DETAILED ERROR LOG', level=1)
    if data.get("errors"):
        for err in data['errors']:
            p = doc.add_paragraph(style='List Bullet')
            runner = p.add_run(f"[{err['category']} - {err['type']}]: ")
            runner.bold = True
            runner.font.color.rgb = RGBColor(200, 0, 0)
            p.add_run(f" '{err['original']}' → '{err['correction']}'")
            p.add_run(f"\n   Reason: {err['explanation']}")
    else:
        doc.add_paragraph("No specific errors detected.")

    # APPENDIX
    doc.add_page_break()
    doc.add_heading('APPENDIX', level=1)
    doc.add_heading('A. Task Prompt:', level=2)
    doc.add_paragraph(topic)
    doc.add_heading('B. Original Essay:', level=2)
    doc.add_paragraph(original_essay)
    doc.add_heading('C. Annotated Version:', level=2)
    clean_annotated = re.sub(r'<[^>]+>', '', data.get("annotatedEssay", "") or "")
    doc.add_paragraph(clean_annotated)

    # D. PROJECTED SCORE
    doc.add_heading('D. PROJECTED BAND SCORE (AFTER REVISION)', level=2)
    rev_scores = data.get("revisedScore")
    if rev_scores:
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        vals = [
            str(rev_scores.get('task_achievement', '-')),
            str(rev_scores.get('cohesion_coherence', '-')),
            str(rev_scores.get('lexical_resource', '-')),
            str(rev_scores.get('grammatical_range', '-')),
            str(rev_scores.get('overall', '-'))
        ]
        for i, h in enumerate(['Task Achievement', 'Coherence', 'Lexical Resource', 'Grammar', 'OVERALL']):
            cell = table.cell(0, i)
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 100, 0)
        for i, v in enumerate(vals):
            table.cell(1, i).text = v
            
        if rev_scores.get('logic_re_evaluation'):
            p = doc.add_paragraph()
            run = p.add_run(f"\nExaminer's Note: {rev_scores['logic_re_evaluation']}")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0, 128, 0)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# PDF EXPORT
def create_pdf(data, topic, original_essay, analysis_text):
    has_font = register_fonts()
    font_name = 'Roboto' if has_font else 'Helvetica'
    font_bold = 'Roboto-Bold' if has_font else 'Helvetica-Bold'

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    
    styles['Title'].fontName = font_name
    styles['Title'].fontSize = 18
    styles['Heading1'].fontName = font_bold
    styles['Heading2'].fontName = font_bold
    styles['Normal'].fontName = font_name
    styles['Normal'].fontSize = 13
    
    h1_style = styles['Heading1']
    h2_style = styles['Heading2']
    normal_style = styles['Normal']
    
    elements = []

    elements.append(Paragraph("IELTS WRITING ASSESSMENT REPORT", styles['Title']))
    elements.append(Spacer(1, 12))

    # 1. BAND SCORE
    elements.append(Paragraph("1. BAND SCORE", h1_style))
    scores = data.get("originalScore")
    
    if scores and isinstance(scores, dict) and scores.get('overall', '-') != '-':
        data_table = [
            ['TA', 'CC', 'LR', 'GRA', 'OVERALL'],
            [
                str(scores.get('task_achievement', '-')),
                str(scores.get('cohesion_coherence', '-')),
                str(scores.get('lexical_resource', '-')),
                str(scores.get('grammatical_range', '-')),
                str(scores.get('overall', '-'))
            ]
        ]
        t = Table(data_table, colWidths=[60, 60, 60, 60, 80])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkred),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name) 
        ]))
        elements.append(t)
    else:
        elements.append(Paragraph("Original score data not found.", normal_style))
    
    elements.append(Spacer(1, 12))

    # 2. ANALYSIS
    elements.append(Paragraph("2. DETAILED ANALYSIS", h1_style))
    if analysis_text:
        safe_text = html.escape(analysis_text).replace('\n', '<br/>').replace('**', '').replace('###', '')
        elements.append(Paragraph(safe_text, normal_style))
    else:
        elements.append(Paragraph("No detailed analysis available.", normal_style))
    elements.append(Spacer(1, 12))

    # 3. ERRORS
    elements.append(Paragraph("3. ERROR LOG", h1_style))
    if data.get("errors"):
        for err in data['errors']:
            cat = html.escape(str(err.get('category', '')))
            typ = html.escape(str(err.get('type', '')))
            orig = html.escape(str(err.get('original', '')))
            fix = html.escape(str(err.get('correction', '')))
            text = f"<b>[{cat}] {typ}</b><br/>Original: <strike>{orig}</strike> -> Fix: <b>{fix}</b>"
            elements.append(Paragraph(text, normal_style))
            elements.append(Spacer(1, 6))

    # APPENDIX
    elements.append(PageBreak())
    elements.append(Paragraph("APPENDIX", h1_style))
    
    elements.append(Paragraph("<b>A. Task Prompt:</b>", h2_style))
    elements.append(Paragraph(html.escape(topic).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>B. Original Essay:</b>", h2_style))
    elements.append(Paragraph(html.escape(original_essay).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("<b>C. Annotated Version:</b>", h2_style))
    clean_annotated = re.sub(r'<[^>]+>', '', data.get("annotatedEssay", "") or "")
    elements.append(Paragraph(html.escape(clean_annotated).replace('\n', '<br/>'), normal_style))
    elements.append(Spacer(1, 10))

    # D. PROJECTED
    elements.append(Paragraph("<b>D. PROJECTED BAND SCORE (AFTER REVISION):</b>", h2_style))
    rev_scores = data.get("revisedScore")
    if rev_scores:
        rev_table_data = [
            ['TA', 'CC', 'LR', 'GRA', 'OVERALL'],
            [
                str(rev_scores.get('task_achievement', '-')),
                str(rev_scores.get('cohesion_coherence', '-')),
                str(rev_scores.get('lexical_resource', '-')),
                str(rev_scores.get('grammatical_range', '-')),
                str(rev_scores.get('overall', '-'))
            ]
        ]
        t2 = Table(rev_table_data, colWidths=[60, 60, 60, 60, 80])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name)
        ]))
        elements.append(t2)
        
        if rev_scores.get('logic_re_evaluation'):
            safe_note = html.escape(rev_scores['logic_re_evaluation'])
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(f"<i>Examiner's Note: {safe_note}</i>", normal_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer
    
# ==========================================
# 4. MAIN UI
# ==========================================

# HEADER
c1, c2 = st.columns([3, 1])
with c1:
    st.markdown("""
        <div style="display: flex; flex-direction: column; justify-content: center;">
            <h1 style="margin-bottom: 5px; line-height: 0.2;">
                IELTS Examiner <span class='pro-badge'>Pro</span>
            </h1>
            <div>
                <span class='verified-badge' style="margin-left: 2px;">
                    🛡️ BC CERTIFIED EXPERT
                </span>
            </div>
        </div>
    """, unsafe_allow_html=True)
with c2:
    if st.button("🗑️ Clear Session", use_container_width=True):
        st.session_state.messages = []
        st.session_state.submitted = False 
        st.rerun()

if "submitted" not in st.session_state:
    st.session_state.submitted = False

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "ai",
            "content": """
<div style="font-family: 'Inter', sans-serif; color: #1F2937; line-height: 1.4; font-size: 16px; max-width: 850px;">
    <h3 style="color: #D40E14; font-family: 'Merriweather', serif; margin-top: 0; margin-bottom: 15px; font-size: 22px; border-bottom: 3px solid #D40E14; display: inline-block; padding-bottom: 5px;">
        Welcome to the Official Task 1 Assessment Room.
    </h3>
    <p style="margin-bottom: 13px;">
        This system provides expert-level evaluation of <b>IELTS Academic Task 1 reports</b>, based on the official IELTS band descriptors.
    </p>
    <p style="margin-bottom: 13px;">
        The assessment focuses on objective, criteria-based feedback to help you understand your current writing level and areas for improvement.
    </p>
    <div style="background-color: #F8FAFC; border-radius: 8px; padding: 15px 20px; border-left: 5px solid #D40E14; box-shadow: 0 2px 4px rgba(0,0,0,0.05);">
        <p style="margin-bottom: 10px; font-weight: 800; font-size: 13px; color: #111827; text-transform: uppercase; letter-spacing: 1px;">
            Guidelines for a valid submission:
        </p>
        <div style="color: #374151;">
            <div style="margin-bottom: 6px;">• <b>Task Prompt:</b> Provide the original question or instruction.</div>
            <div style="margin-bottom: 6px;">• <b>Visual Data:</b> Upload a clear image of the chart, graph, table, or diagram.</div>
            <div>• <b>Written Report:</b> Paste your complete response (at least <b>150 words</b> to avoid penalties).</div>
        </div>
    </div>
</div>
""",
            "data": None
        }
    ]

for msg in st.session_state.messages:
    with st.chat_message(msg["role"], avatar="👨‍🏫" if msg["role"] == "ai" else "👤"):
        if msg["role"] == "user":
            if msg.get("topic"):
                st.markdown(f"**📝 Task Prompt:**\n> {msg['topic']}")
            if msg.get("image"):
                st.image(msg["image"], caption="Visual Resource Attached", width=400)
            st.write(msg["content"])
        else:
            st.markdown(f'<div class="report-content">{msg["content"]}</div>', unsafe_allow_html=True)  
            
            if msg.get("data") and msg["data"]["errors"]:
                all_errors = msg["data"]["errors"]
                # Filter errors (Assuming AI outputs categories in English now)
                micro_errors = [e for e in all_errors if e.get('category') in ['Grammar', 'Vocabulary']]
                macro_errors = [e for e in all_errors if e.get('category') not in ['Grammar', 'Vocabulary']]

                # --- 1. GRAMMAR & VOCAB ---
                if micro_errors:
                    with st.expander(f"🚩 Grammar & Vocabulary Corrections ({len(micro_errors)} Issues)", expanded=True):
                        for idx, err in enumerate(micro_errors):
                            cat = err.get('category', 'Grammar')
                            badge_style = "background:#DCFCE7; color:#166534; border:1px solid #86EFAC" if cat == 'Grammar' else "background:#FEF9C3; color:#854D0E; border:1px solid #FCD34D"
                            impact = err.get('impact_level', 'Low').upper()
                            
                            html_micro = textwrap.dedent(f"""
                                <div class="error-card" style="margin-bottom:12px; border: 1px solid #eee; padding: 10px; border-radius: 8px;">
                                    <div style="display:flex; justify-content:space-between; align-items:center; border-bottom:1px solid #eee; padding-bottom:4px; margin-bottom:4px;">
                                        <div style="display:flex; align-items:center;">
                                            <span style="background:#F3F4F6; width:22px; height:22px; display:inline-flex; align-items:center; justify-content:center; border-radius:50%; font-weight:bold; font-size:12px; margin-right:8px;">{idx + 1}</span>
                                            <span style="{badge_style}; padding: 2px 8px; border-radius: 6px; font-size: 11px; font-weight: 800; text-transform: uppercase;">{cat}</span>
                                            <span style="font-weight:700; font-size:16px; margin-left:10px; color:#1F2937;">{err['type']}</span>
                                        </div>
                                        <span style="background:#F3F4F6; color:#666; padding:2px 8px; border-radius:6px; font-size:10px; font-weight:bold;">{impact}</span>
                                    </div>
                                    <div style="background:#F9FAFB; padding:10px; border-radius:6px; font-size:15px; line-height: 1.5;">
                                        <div style="margin-bottom:4px;">
                                            <span style="color:#6B7280; font-size:14px; font-weight:800; letter-spacing: 0.5px;">ORIGINAL:</span> 
                                            <span style="text-decoration:line-through; color:#9CA3AF; margin-left: 6px;">{err['original']}</span>
                                        </div>
                                        <div>
                                            <span style="color:#6B7280; font-size:14px; font-weight:800; letter-spacing: 0.5px;">FIX:</span> 
                                            <span style="{badge_style}; padding:1px 6px; border-radius:4px; font-weight:bold; margin-left: 6px; color:#111;">{err['correction']}</span>
                                        </div>
                                    </div>
                                    <div style="font-size:14px; color:#4B5563; margin-top:6px; font-style: italic;">
                                        Note: {err['explanation']}
                                    </div>
                                </div>
                            """).strip()
                            st.markdown(html_micro, unsafe_allow_html=True)

                # --- 2. COHERENCE & COHESION ---
                if macro_errors:
                    st.markdown("---") 
                    st.markdown(f"#### 💡 Coherence & Cohesion Improvements ({len(macro_errors)} Issues)")
                    st.caption("Focus on logical flow, grouping, and data representation.")
                    
                    with st.expander("View Logic & Coherence Details", expanded=True):
                        for idx, err in enumerate(macro_errors):
                            badge_style = "background:#DBEAFE; color:#1E40AF; border:1px solid #BFDBFE"
                            impact = str(err.get('impact_level', 'Low')).upper()
                            err_type = str(err.get('type', 'Logic Error'))
                            explanation = str(err.get('explanation', ''))
                            original = str(err.get('original', ''))
                            correction = str(err.get('correction', ''))

                            html_macro = """
<div class="error-card" style="border-left: 5px solid #3B82F6; margin-bottom:16px; background: white; padding: 16px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); border-top: 1px solid #eee; border-right: 1px solid #eee; border-bottom: 1px solid #eee;">
    <div class="error-header" style="display:flex; justify-content:space-between; align-items:center; border-bottom:1px solid #eee; padding-bottom:8px; margin-bottom:12px;">
        <div style="display:flex; align-items:center;">
            <span style="{badge_style}; padding: 2px 8px; border-radius: 6px; font-size: 11px; font-weight: 800; text-transform: uppercase;">COHERENCE & COHESION</span>
            <span style="font-weight:700; font-size:18px; margin-left:12px; color:#1F2937;">{err_type}</span>
        </div>
        <span style="background:#F3F4F6; color:#666; padding:4px 10px; border-radius:6px; font-size:11px; font-weight:bold;">{impact}</span>
    </div>
    <div style="font-size:16px; color:#374151; line-height: 1.6;">
        <div style="margin-bottom: 8px;">
            <span style="font-weight:800; color:#1E40AF; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px;">ISSUE:</span> 
            <span>{explanation}</span>
        </div>
        <div style="margin-bottom: 8px;">
            <span style="font-weight:800; color:#6B7280; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px;">ORIGINAL:</span> 
            <span style="text-decoration:line-through; color:#9CA3AF;">{original}</span>
        </div>
        <div>
            <span style="font-weight:800; color:#059669; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px;">SUGGESTED FIX:</span> 
            <span style="font-weight:600; color:#111;">{correction}</span>
        </div>
    </div>
</div>
""".format(badge_style=badge_style, err_type=err_type, impact=impact, explanation=explanation, original=original, correction=correction)

                            st.markdown(html_macro, unsafe_allow_html=True)
                else:
                    structure_breakers = ['Fragment', 'Run-on Sentence', 'Comma Splice', 'Sentence Structure']
                    has_structure_error = any(e.get('type') in structure_breakers for e in all_errors)
                    st.markdown("---")
                    st.markdown("#### 💡 Coherence & Cohesion Review")
                    if has_structure_error:
                        st.warning("⚠️ **Note:** Although there are no major logic errors, structural errors in the Grammar section above are negatively affecting coherence.")
                    else:
                        st.success("✅ **Excellent!** The essay has a coherent structure and ideas are well-linked.")

            # 3. Annotated Essay
            if msg.get("data") and msg["data"]["annotatedEssay"]:
                st.markdown("### 📝 Examiner's Annotated Report")
                st.caption("The essay has been corrected (strikethrough = incorrect, highlighted = corrected)")
                st.markdown(f'<div class="annotated-text">{msg["data"]["annotatedEssay"]}</div>', unsafe_allow_html=True)
            
            # 4. Revised Score
            if msg.get("data") and msg["data"].get("revisedScore"):
                scores = msg["data"]["revisedScore"]
                
                st.markdown("### 📊 Projected Band (Revised Version)")
                
                if float(str(scores.get('overall', 0)).replace('-', '0')) >= 8.5:
                    st.success("✨ This revised version is approaching perfection.")
                else:
                    st.warning(f"⚠️ **Examiner's Note:** This revised version only reached {scores.get('overall')} because: {scores.get('logic_re_evaluation', 'it still lacks the absolute conciseness of Band 9.0')}")

                cols = st.columns(5)
                cols[0].metric("TA", scores.get("task_achievement", "-"))
                cols[1].metric("CC", scores.get("cohesion_coherence", "-"))
                cols[2].metric("LR", scores.get("lexical_resource", "-"))
                cols[3].metric("GRA", scores.get("grammatical_range", "-"))
                cols[4].metric("OVERALL", scores.get("overall", "-"))
                
                # --- DOWNLOAD BUTTONS ---
                st.markdown("---")
                st.markdown("### 📥 Download Report")
                
                topic_text = msg.get("topic", "")
                essay_text = msg.get("original_essay", "")
                analysis_text = msg.get("content", "")
                
                if not topic_text:
                    try:
                        prev_msg_index = st.session_state.messages.index(msg) - 1
                        if prev_msg_index >= 0:
                            prev_msg = st.session_state.messages[prev_msg_index]
                            topic_text = prev_msg.get("topic", "Topic not found")
                            essay_text = prev_msg.get("content", "Essay not found")
                    except:
                        pass

                d1, d2 = st.columns(2)
                
                docx_file = create_docx(msg["data"], topic_text, essay_text, analysis_text)
                d1.download_button(
                    label="📄 Download Analysis (.docx)",
                    data=docx_file,
                    file_name=f"IELTS_Report_{int(time.time())}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                pdf_file = create_pdf(msg["data"], topic_text, essay_text, analysis_text)
                d2.download_button(
                    label="📕 Download Analysis (.pdf)",
                    data=pdf_file,
                    file_name=f"IELTS_Report_{int(time.time())}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

# ==========================================
# 5. INPUT AREA
# ==========================================

if not st.session_state.submitted:
    st.markdown("---")
    with st.container():
        col_left, col_right = st.columns([1.3, 2.7], gap="large")
        
        with col_left:
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>📝 TASK 1 QUESTION / PROMPT</p>", unsafe_allow_html=True)
            topic_input = st.text_area("topic_label", label_visibility="collapsed", height=280, placeholder="Paste the official question text here...")
            
            st.markdown("<div style='margin-top: 25px;'></div>", unsafe_allow_html=True)
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>📊 VISUAL DATA</p>", unsafe_allow_html=True)
            uploaded_file = st.file_uploader("file_label", label_visibility="collapsed", type=['png', 'jpg', 'jpeg'])
            
        with col_right:
            st.markdown("<p style='font-weight: 700; font-size: 15px; color: #1F2937;'>✍️ YOUR WRITTEN REPORT</p>", unsafe_allow_html=True)
            essay_input = st.text_area("essay_label", label_visibility="collapsed", height=515, placeholder="Type or paste your response here (aim for 150+ words)...")

        st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)
        submit_btn = st.button("🚀 SUBMIT FOR ASSESSMENT", type="primary", use_container_width=True)

        if submit_btn:
            # VALIDATION
            if not topic_input.strip():
                st.warning("⚠️ Required: Please enter the Task Prompt!")
            elif uploaded_file is None:
                st.warning("⚠️ Required: Please upload the Visual Data (Chart/Graph)!")
            elif not essay_input.strip() or len(essay_input.strip()) < 10:
                st.warning("⚠️ Required: Please enter your essay (at least 10 characters)!")
            else:
                # LOADING SEQUENCE
                loading_steps = [
                    "🕵️ INITIAL VALIDATION: IDENTIFYING EXAM CONTEXT AND ENFORCING WORD COUNT CONSTRAINTS...",
                    "🔍 EXHAUSTIVE ERROR SCANNING: CONDUCTING SENTENCE-BY-SENTENCE REVIEW FOR ALL ERRORS...",
                    "📊 DEEP CRITERIA ANALYSIS: EVALUATING TA, CC, LR, AND GRA STANDARDS WITH CEILING SCORES...",
                    "🧮 SCORE CALCULATION: DETERMINING COMPONENT BANDS AND APPLYING IELTS ROUNDING RULES...",
                    "⚖️ CONSISTENCY CHECK: CROSS-REFERENCING ASSIGNED SCORES WITH ERROR LOG FOR LOGICAL ACCURACY...",
                    "📝 OUTPUT GENERATION: COMPILING DETAILED ANALYSIS AND PRODUCING ANNOTATED ESSAY DATA..."
                ]
                
                status_container = st.status("👨‍🏫 Senior Examiner is starting assessment...", expanded=True)
                progress_bar = status_container.progress(0)
                
                try:
                    # 1. Process Image
                    image_part = Image.open(uploaded_file)
                    
                    # 2. Prepare Prompt
                    full_prompt = GRADING_PROMPT_TEMPLATE.replace('{{TOPIC}}', topic_input).replace('{{ESSAY}}', essay_input)
                    
                    # 3. Call AI
                    response, used_model = generate_content_with_failover(full_prompt, image_part)
                    
                    # 4. Loading Animation
                    for i, text in enumerate(loading_steps):
                        status_container.write(text)
                        progress_bar.progress(int((i + 1) * (100 / len(loading_steps))))
                        time.sleep(2.8) 
                    
                    if response:
                        markdown_text, parsed_data = process_response(response.text)
                        st.session_state.messages.append({"role": "user", "content": essay_input, "topic": topic_input, "image": uploaded_file})
                        st.session_state.messages.append({"role": "ai", "content": markdown_text, "data": parsed_data, "model_version": used_model})
                        st.session_state.submitted = True
                        status_container.update(label="✅ ASSESSMENT COMPLETE!", state="complete", expanded=False)
                        st.rerun()
                        
                except Exception as e:
                    status_container.update(label="❌ Error occurred!", state="error")
                    st.error(f"System Error: {e}")

# Footer
st.markdown("---")

st.caption("Developed by Albert Nguyen - v20251225.")


